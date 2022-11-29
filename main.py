from docx import Document
from docx.shared import Inches
import xlsxwriter

def main():
    guestnames = ['Айгерим', 'Алмаз', 'Артур', 'Дмитрий и Виктория']
    for x in guestnames:
        document = Document()
        document.add_heading('Приглашение на Новогодний Корпоратив', 0)
        p = document.add_paragraph(x + ', спешим пригласить Вас на Новогодний Корпоратив в ')
        p.add_run(' Lounge Bar "Миндаль"').bold = True
        p.add_run(', который состоится 24.12.2022 на ул. Абая 166')
        document.add_heading('Мы ждём вас, на наш праздник!!!', level=1)
        document.add_paragraph('')
        document.add_picture('1.png', width=Inches(5.25))
        document.add_page_break()
        document.save(x + '.docx')

    try:
        my_file = 'demo.xlsx'  # Имя файла
        book = xlsxwriter.Workbook(my_file)  # Создание файла
        sheet = book.add_worksheet()  # Добавление в него книги
        sheet.set_column('A:A', 20)   # Установка ширины колонки
        bold = book.add_format({'bold': True})  # Формат жирного текста
        sheet.write('A1', 'Приглашение')  # Выдача текста в ячейку
        sheet.write('A2', 'на Новогодний Корпоратив')  # Выдача жирного текста в ячейку
        sheet.write(2, 1, 'Lounge Bar "Миндаль"', bold)  # Выдача значения в ячейку 3 строка 1 столбец [2,0]
        sheet.write(3, 3, '24.12.2022')   # Выдача значения в ячейку 4 строка 2 столбец [3,1]
        sheet.insert_image('B5', '1.png')  # Вставка в ячейку картинки
        book.close()  # Закрытие файла
    except Exception as a:  # Обработка ошибок
        print("Error!")
        print(a)

if __name__ == '__main__':
    main()


